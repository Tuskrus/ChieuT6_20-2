from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from datetime import datetime

# Create a new Document
doc = Document()

# Set document title and formatting
title = doc.add_paragraph()
title_run = title.add_run('Inventory Management System')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add timestamp
timestamp = doc.add_paragraph()
timestamp_run = timestamp.add_run(f'Document Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
timestamp_run.font.size = Pt(10)
timestamp_run.font.italic = True
timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # spacing

# 1. Project Overview
doc.add_heading('1. Project Overview', level=1)
overview = doc.add_paragraph(
    'This document describes the Inventory Management System implementation for the NNPTUD-C6 Node.js project. '
    'The system manages product inventory including stock, reservations, and sold items.'
)

# 2. Database Schema
doc.add_heading('2. Database Schema', level=1)

doc.add_heading('Inventory Model', level=2)
schema_text = doc.add_paragraph()
schema_text.add_run('The Inventory schema contains the following fields:\n\n').bold = True

properties = [
    ('product', 'ObjectID (ref: product)', 'Required, Unique - Reference to Product'),
    ('stock', 'Number', 'Default: 0, Min: 0 - Available stock quantity'),
    ('reserved', 'Number', 'Default: 0, Min: 0 - Reserved quantity'),
    ('soldCount', 'Number', 'Default: 0, Min: 0 - Total number of items sold'),
    ('timestamps', 'createdAt, updatedAt', 'Automatic tracking of creation and update times')
]

table = doc.add_table(rows=len(properties) + 1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Field'
hdr_cells[1].text = 'Type'
hdr_cells[2].text = 'Description'

for i, (field, ftype, description) in enumerate(properties, 1):
    cells = table.rows[i].cells
    cells[0].text = field
    cells[1].text = ftype
    cells[2].text = description

# 3. API Endpoints
doc.add_heading('3. API Endpoints', level=1)

endpoints = [
    {
        'method': 'GET',
        'endpoint': '/api/v1/inventories',
        'description': 'Get all inventories with product details',
        'params': 'None',
        'example_response': {
            'success': True,
            'data': [
                {
                    '_id': '65a1b2c3d4e5f6g7h8i9j0k1',
                    'product': {
                        '_id': '65a1b2c3d4e5f6g7h8i9j0k0',
                        'title': 'Product 1',
                        'price': 100,
                        'category': '65a1b2c3d4e5f6g7h8i9j0k2'
                    },
                    'stock': 50,
                    'reserved': 10,
                    'soldCount': 100,
                    'createdAt': '2024-01-15T10:00:00.000Z',
                    'updatedAt': '2024-01-15T10:00:00.000Z'
                }
            ]
        }
    },
    {
        'method': 'GET',
        'endpoint': '/api/v1/inventories/:id',
        'description': 'Get inventory by ID with full product details',
        'params': 'id (inventory ID in URL)',
        'example_response': {
            'success': True,
            'data': {
                '_id': '65a1b2c3d4e5f6g7h8i9j0k1',
                'product': {
                    '_id': '65a1b2c3d4e5f6g7h8i9j0k0',
                    'title': 'Product 1',
                    'price': 100,
                    'description': 'Product description',
                    'category': '65a1b2c3d4e5f6g7h8i9j0k2',
                    'images': ['url1', 'url2']
                },
                'stock': 50,
                'reserved': 10,
                'soldCount': 100
            }
        }
    },
    {
        'method': 'POST',
        'endpoint': '/api/v1/inventories/add-stock',
        'description': 'Increase stock for a product',
        'params': 'Body: { product: ObjectID, quantity: number }',
        'example_response': {
            'success': True,
            'message': 'Stock added successfully',
            'data': {
                'stock': 75,
                'reserved': 10,
                'soldCount': 100
            }
        }
    },
    {
        'method': 'POST',
        'endpoint': '/api/v1/inventories/remove-stock',
        'description': 'Decrease stock for a product',
        'params': 'Body: { product: ObjectID, quantity: number }',
        'example_response': {
            'success': True,
            'message': 'Stock removed successfully',
            'data': {
                'stock': 25,
                'reserved': 10,
                'soldCount': 100
            }
        }
    },
    {
        'method': 'POST',
        'endpoint': '/api/v1/inventories/reservation',
        'description': 'Reserve stock (decreases stock, increases reserved)',
        'params': 'Body: { product: ObjectID, quantity: number }',
        'example_response': {
            'success': True,
            'message': 'Stock reserved successfully',
            'data': {
                'stock': 40,
                'reserved': 20,
                'soldCount': 100
            }
        }
    },
    {
        'method': 'POST',
        'endpoint': '/api/v1/inventories/sold',
        'description': 'Mark reserved items as sold (decreases reserved, increases soldCount)',
        'params': 'Body: { product: ObjectID, quantity: number }',
        'example_response': {
            'success': True,
            'message': 'Inventory marked as sold successfully',
            'data': {
                'stock': 40,
                'reserved': 10,
                'soldCount': 110
            }
        }
    }
]

for endpoint in endpoints:
    doc.add_heading(f"{endpoint['method']} {endpoint['endpoint']}", level=2)
    
    doc.add_paragraph(f"Description: {endpoint['description']}")
    doc.add_paragraph(f"Parameters: {endpoint['params']}")
    
    doc.add_paragraph('Example Response:', style='List Bullet')
    response_para = doc.add_paragraph()
    response_run = response_para.add_run(json.dumps(endpoint['example_response'], indent=2))
    response_run.font.name = 'Courier New'
    response_run.font.size = Pt(9)

# 4. Testing Instructions
doc.add_heading('4. Testing Instructions', level=1)

doc.add_heading('Prerequisites', level=2)
prereq = doc.add_paragraph(style='List Bullet')
prereq.add_run('MongoDB running on localhost:27017')
prereq = doc.add_paragraph(style='List Bullet')
prereq.add_run('Node.js dependencies installed (npm install)')
prereq = doc.add_paragraph(style='List Bullet')
prereq.add_run('Postman or similar API testing tool')
prereq = doc.add_paragraph(style='List Bullet')
prereq.add_run('Base URL: http://localhost:3000')

doc.add_heading('Step 1: Start MongoDB', level=2)
doc.add_paragraph('Start MongoDB service on your system. The application expects MongoDB to be running on localhost:27017.')

doc.add_heading('Step 2: Start the Application', level=2)
code = doc.add_paragraph('npm start')
code.style = 'Heading 3'
code_para = doc.add_paragraph()
code_run = code_para.add_run('The application will start on port 3000')
code_run.font.italic = True

doc.add_heading('Step 3: Test Create Product', level=2)
doc.add_paragraph('First, create a product. This will automatically create an inventory record.')
doc.add_paragraph('Method: POST', style='List Bullet')
doc.add_paragraph('Endpoint: /api/v1/products', style='List Bullet')
doc.add_paragraph('Body:', style='List Bullet')

body_code = doc.add_paragraph()
body_run = body_code.add_run('''
{
  "title": "Test Product",
  "price": 99.99,
  "description": "A test product",
  "categoryId": "CATEGORY_ID_HERE",
  "images": ["image_url"]
}
''')
body_run.font.name = 'Courier New'
body_run.font.size = Pt(9)

test_steps = [
    ('Get All Inventories', 'GET', '/api/v1/inventories', 'None', 'Returns list of all inventories'),
    ('Get Inventory by ID', 'GET', '/api/v1/inventories/:id', 'Inventory ID from previous response', 'Returns single inventory with product details'),
    ('Add Stock', 'POST', '/api/v1/inventories/add-stock', '{"product": "PRODUCT_ID", "quantity": 50}', 'Increases stock by 50'),
    ('Reserve Stock', 'POST', '/api/v1/inventories/reservation', '{"product": "PRODUCT_ID", "quantity": 20}', 'Decreases stock by 20, increases reserved by 20'),
    ('Mark as Sold', 'POST', '/api/v1/inventories/sold', '{"product": "PRODUCT_ID", "quantity": 10}', 'Decreases reserved by 10, increases soldCount by 10'),
    ('Remove Stock', 'POST', '/api/v1/inventories/remove-stock', '{"product": "PRODUCT_ID", "quantity": 5}', 'Decreases stock by 5'),
]

doc.add_heading('Test Cases', level=2)
test_table = doc.add_table(rows=len(test_steps) + 1, cols=5)
test_table.style = 'Light Grid Accent 1'
hdr = test_table.rows[0].cells
hdr[0].text = 'Test Case'
hdr[1].text = 'Method'
hdr[2].text = 'Endpoint'
hdr[3].text = 'Request Body'
hdr[4].text = 'Expected Result'

for i, (name, method, endpoint, body, result) in enumerate(test_steps, 1):
    cells = test_table.rows[i].cells
    cells[0].text = name
    cells[1].text = method
    cells[2].text = endpoint
    cells[3].text = body
    cells[4].text = result

# 5. Implementation Details
doc.add_heading('5. Implementation Details', level=1)

doc.add_heading('File Structure', level=2)
files = [
    ('schemas/inventory.js', 'MongoDB schema definition for inventory'),
    ('controllers/inventory.js', 'Business logic for inventory operations'),
    ('routes/inventory.js', 'API endpoint definitions'),
    ('routes/products.js', 'Updated to auto-create inventory'),
    ('app.js', 'Updated to register inventory routes'),
]

file_table = doc.add_table(rows=len(files) + 1, cols=2)
file_table.style = 'Light Grid Accent 1'
hdr = file_table.rows[0].cells
hdr[0].text = 'File'
hdr[1].text = 'Purpose'

for i, (fname, purpose) in enumerate(files, 1):
    cells = file_table.rows[i].cells
    cells[0].text = fname
    cells[1].text = purpose

# 6. Git Commits
doc.add_heading('6. Version Control', level=1)

doc.add_heading('Commits Made', level=2)
commit_info = doc.add_paragraph()
commit_run = commit_info.add_run(
    'Commit: b4cb629\n'
    'Branch: 20260313\n'
    'Message: feat: implement inventory management system\n\n'
    'Changes:\n'
    '• schemas/inventory.js - New inventory schema\n'
    '• controllers/inventory.js - Inventory business logic\n'
    '• routes/inventory.js - Inventory API endpoints\n'
    '• routes/products.js - Auto-create inventory on product creation\n'
    '• app.js - Register inventory routes'
)
commit_run.font.name = 'Courier New'
commit_run.font.size = Pt(9)

# 7. Error Handling
doc.add_heading('7. Error Handling', level=1)

errors = [
    ('400 Bad Request', 'Missing required fields or invalid data'),
    ('400 Bad Request', 'Insufficient stock to remove or reserve'),
    ('400 Bad Request', 'Insufficient reserved stock to mark as sold'),
    ('404 Not Found', 'Inventory or product not found'),
    ('500 Internal Server Error', 'Database or server error'),
]

error_table = doc.add_table(rows=len(errors) + 1, cols=2)
error_table.style = 'Light Grid Accent 1'
hdr = error_table.rows[0].cells
hdr[0].text = 'Error Code'
hdr[1].text = 'Description'

for i, (code, desc) in enumerate(errors, 1):
    cells = error_table.rows[i].cells
    cells[0].text = code
    cells[1].text = desc

# Save the document
output_path = 'c:\\Users\\LENOVO\\OneDrive\\Desktop\\20-3\\NNPTUD-C6\\Inventory_Management_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
